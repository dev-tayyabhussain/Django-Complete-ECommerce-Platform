from docx import Document

# Paste your full markdown text here as a raw string
text = r"""
# E-Commerce DevOps Implementation: Academic Case Study

## Executive Summary

This academic case study presents a comprehensive implementation of modern DevOps practices...
"""

# Create a Word document
doc = Document()

# Split the text into lines for parsing
lines = text.splitlines()
inside_code_block = False
code_buffer = []

for line in lines:
    if line.startswith("### "):
        doc.add_heading(line.replace("### ", ""), level=3)
    elif line.startswith("## "):
        doc.add_heading(line.replace("## ", ""), level=2)
    elif line.startswith("# "):
        doc.add_heading(line.replace("# ", ""), level=1)
    elif line.startswith("```"):
        if inside_code_block:
            # End of code block
            doc.add_paragraph("\n".join(code_buffer), style="Normal")
            code_buffer = []
            inside_code_block = False
        else:
            # Start of code block
            inside_code_block = True
    else:
        if inside_code_block:
            code_buffer.append(line)
        else:
            doc.add_paragraph(line)

# Save document
output_path = "ACADEMIC_CASE_STUDY_FULL.docx"
doc.save(output_path)
print(f"Saved as {output_path}")
